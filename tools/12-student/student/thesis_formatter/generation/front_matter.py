"""Generate front matter pages: title page, declaration, abstract."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_title_page(doc_path: str, project: dict, profile: dict) -> str:
    output_path = _output_path(doc_path, "_title")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    body = doc.element.body

    title_paras = [
        _centered_para(doc, project.get("university", ""), Pt(14), bold=True),
        _centered_para(doc, "", Pt(12)),
        _centered_para(doc, project.get("title", ""), Pt(16), bold=True),
        _centered_para(doc, "", Pt(12)),
        _centered_para(doc, f"by {project.get('author', '')}", Pt(14)),
        _centered_para(doc, "", Pt(12)),
        _centered_para(doc, f"A thesis submitted in partial fulfilment of the requirements", Pt(12)),
        _centered_para(doc, f"for the degree of {project.get('degree', '')}", Pt(12)),
        _centered_para(doc, f"in the {project.get('department', '')}", Pt(12)),
        _centered_para(doc, "", Pt(12)),
        _centered_para(doc, f"Supervisor: {project.get('supervisor', '')}", Pt(12)),
        _centered_para(doc, str(project.get("year", "")), Pt(12)),
    ]

    for i, para in enumerate(title_paras):
        body.insert(i, para._element)

    doc.save(output_path)
    return output_path


def generate_declaration(doc_path: str, university: str) -> str:
    output_path = _output_path(doc_path, "_decl")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    body = doc.element.body

    heading = doc.add_paragraph("Declaration", style="Heading 1")
    body.append(heading._element)

    text = (
        f"I hereby declare that this thesis represents my own work which has been done "
        f"after registration for the degree at {university}, and has not been previously "
        f"included in a thesis or dissertation submitted to this or any other institution "
        f"for a degree, diploma, or other qualifications.\n\n"
        f"I have read the University's current research ethics guidelines, and accept "
        f"responsibility for the conduct of the procedures in accordance with the "
        f"University's Research Ethics Committee."
    )
    decl_para = doc.add_paragraph(text)
    body.append(decl_para._element)

    sig_para = doc.add_paragraph("\n\nSigned: _________________________\n\nDate: _________________________")
    body.append(sig_para._element)

    doc.save(output_path)
    return output_path


def generate_abstract_page(doc_path: str, abstract_en: str, abstract_tc: str) -> str:
    output_path = _output_path(doc_path, "_abstract")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    body = doc.element.body

    heading_en = doc.add_paragraph("Abstract", style="Heading 1")
    body.append(heading_en._element)
    en_para = doc.add_paragraph(abstract_en)
    body.append(en_para._element)

    if abstract_tc:
        heading_tc = doc.add_paragraph("摘要", style="Heading 1")
        body.append(heading_tc._element)
        tc_para = doc.add_paragraph(abstract_tc)
        body.append(tc_para._element)

    doc.save(output_path)
    return output_path


def _centered_para(doc: Document, text: str, font_size: Pt, bold: bool = False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = font_size
    run.font.bold = bold
    return para


def _output_path(doc_path: str, suffix: str = "") -> str:
    p = Path(doc_path)
    return str(p.parent / f"{p.stem}{suffix}{p.suffix}")
