"""Generate Table of Contents from document headings."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def generate_toc(doc_path: str) -> str:
    output_path = _output_path(doc_path, "_toc")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    toc_para = _find_or_insert_toc_paragraph(doc)
    _insert_toc_field(toc_para)
    doc.save(output_path)
    return output_path


def _find_or_insert_toc_paragraph(doc: Document):
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() in ("table of contents", "contents"):
            if i + 1 < len(doc.paragraphs):
                return doc.paragraphs[i + 1]
            return doc.add_paragraph()

    toc_heading = doc.add_paragraph("Table of Contents", style="Heading 1")
    toc_para = doc.add_paragraph()

    body = doc.element.body
    body.insert(2, toc_heading._element)
    body.insert(3, toc_para._element)
    return toc_para


def _insert_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r' TOC \o "1-3" \h \z \u '
    instr_run._element.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    sep_run = paragraph.add_run()
    sep_run._element.append(fld_char_separate)

    placeholder = paragraph.add_run("[Update field to see Table of Contents]")
    placeholder.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    end_run._element.append(fld_char_end)


def _output_path(doc_path: str, suffix: str = "") -> str:
    p = Path(doc_path)
    return str(p.parent / f"{p.stem}{suffix}{p.suffix}")
