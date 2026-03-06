"""Generate List of Tables by scanning document for table captions."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document


def generate_lot(doc_path: str) -> str:
    output_path = _output_path(doc_path, "_lot")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    tables = _scan_tables(doc)

    if not tables:
        return output_path

    lot_heading = doc.add_paragraph("List of Tables", style="Heading 1")
    body = doc.element.body
    insert_pos = _find_insert_position(doc)
    body.insert(insert_pos, lot_heading._element)
    insert_pos += 1

    for tbl in tables:
        entry = doc.add_paragraph(f"{tbl['label']}\t{tbl['caption']}")
        entry.style = doc.styles["Normal"]
        body.insert(insert_pos, entry._element)
        insert_pos += 1

    doc.save(output_path)
    return output_path


def _scan_tables(doc: Document) -> list[dict]:
    tables = []
    tbl_pattern = re.compile(r"^(Table\s+\d+[\.\d]*)[:\.\s]+(.+)", re.IGNORECASE)
    for para in doc.paragraphs:
        match = tbl_pattern.match(para.text.strip())
        if match:
            tables.append({"label": match.group(1), "caption": match.group(2).strip()})
    return tables


def _find_insert_position(doc: Document) -> int:
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if text == "list of figures":
            return doc.element.body.index(para._element) + 2
        if text in ("table of contents", "contents"):
            return doc.element.body.index(para._element) + 2
    return 2


def _output_path(doc_path: str, suffix: str = "") -> str:
    p = Path(doc_path)
    return str(p.parent / f"{p.stem}{suffix}{p.suffix}")
