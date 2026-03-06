"""Verify that all required thesis sections are present."""

from __future__ import annotations

import json
import re

from docx import Document


def check_completeness(doc_path: str, profile: dict) -> list[dict]:
    doc = Document(doc_path)
    required = profile.get("required_sections", [])
    if isinstance(required, str):
        required = json.loads(required)

    headings = [para.text.strip().lower() for para in doc.paragraphs if para.style and para.style.name.startswith("Heading")]
    full_text_lower = " ".join(para.text.strip().lower() for para in doc.paragraphs)

    results: list[dict] = []
    for section in required:
        found = _section_found(section, headings, full_text_lower)
        results.append({
            "check_type": "completeness",
            "passed": found,
            "message": f"Section '{section}': {'found' if found else 'MISSING'}",
            "location": "document structure",
            "severity": "info" if found else "error",
        })

    return results


_SECTION_PATTERNS: dict[str, list[str]] = {
    "cover": ["cover page", "cover"],
    "title": ["title page", "title"],
    "declaration": ["declaration", "statement of originality"],
    "abstract_en": ["abstract"],
    "abstract_tc": ["摘要", "abstract (chinese)", "abstract (tc)"],
    "acknowledgments": ["acknowledgment", "acknowledgement", "acknowledgments", "acknowledgements"],
    "toc": ["table of contents", "contents"],
    "lof": ["list of figures"],
    "lot": ["list of tables"],
    "chapter": ["chapter"],
    "appendix": ["appendix", "appendices"],
    "bibliography": ["bibliography", "references", "works cited"],
}


def _section_found(section: str, headings: list[str], full_text: str) -> bool:
    patterns = _SECTION_PATTERNS.get(section, [section])
    for pattern in patterns:
        for heading in headings:
            if pattern in heading:
                return True
        if pattern in full_text:
            return True
    return False
