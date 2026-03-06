"""Check document formatting against university profile requirements."""

from __future__ import annotations

import json

from docx import Document
from docx.shared import Mm, Pt


def check_format(doc_path: str, profile: dict) -> list[dict]:
    doc = Document(doc_path)
    results: list[dict] = []

    results.extend(_check_margins(doc, profile))
    results.extend(_check_fonts(doc, profile))
    results.extend(_check_spacing(doc, profile))

    return results


def _check_margins(doc: Document, profile: dict) -> list[dict]:
    results = []
    margins = profile.get("margins", {})
    if isinstance(margins, str):
        margins = json.loads(margins)

    expected_top = margins.get("top", profile.get("margin_top", 25))
    expected_bottom = margins.get("bottom", profile.get("margin_bottom", 25))
    expected_left = margins.get("left", profile.get("margin_left", 25))
    expected_right = margins.get("right", profile.get("margin_right", 25))

    for i, section in enumerate(doc.sections):
        loc = f"Section {i + 1}"
        actual_top = round(section.top_margin / Mm(1), 1) if section.top_margin else 0
        actual_bottom = round(section.bottom_margin / Mm(1), 1) if section.bottom_margin else 0
        actual_left = round(section.left_margin / Mm(1), 1) if section.left_margin else 0
        actual_right = round(section.right_margin / Mm(1), 1) if section.right_margin else 0

        for name, actual, expected in [
            ("top margin", actual_top, expected_top),
            ("bottom margin", actual_bottom, expected_bottom),
            ("left margin", actual_left, expected_left),
            ("right margin", actual_right, expected_right),
        ]:
            passed = abs(actual - expected) < 2
            results.append({
                "check_type": "margin",
                "passed": passed,
                "message": f"{name}: {actual}mm (expected {expected}mm)" if not passed else f"{name}: OK",
                "location": loc,
                "severity": "error" if not passed else "info",
            })

    return results


def _check_fonts(doc: Document, profile: dict) -> list[dict]:
    results = []
    expected_name = profile.get("font_name", "Times New Roman")
    expected_size = profile.get("font_size", 12)

    wrong_fonts = set()
    wrong_sizes = set()

    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith("Heading"):
            continue
        for run in para.runs:
            if run.font.name and run.font.name != expected_name:
                wrong_fonts.add((i + 1, run.font.name))
            if run.font.size and run.font.size != Pt(expected_size):
                actual_pt = round(run.font.size / Pt(1), 1)
                wrong_sizes.add((i + 1, actual_pt))

    if wrong_fonts:
        sample = list(wrong_fonts)[:3]
        locations = ", ".join(f"para {p}" for p, _ in sample)
        results.append({
            "check_type": "font",
            "passed": False,
            "message": f"Non-standard font found (expected {expected_name})",
            "location": locations,
            "severity": "warning",
        })
    else:
        results.append({
            "check_type": "font",
            "passed": True,
            "message": f"Font: {expected_name} OK",
            "location": "document",
            "severity": "info",
        })

    if wrong_sizes:
        sample = list(wrong_sizes)[:3]
        locations = ", ".join(f"para {p} ({s}pt)" for p, s in sample)
        results.append({
            "check_type": "font_size",
            "passed": False,
            "message": f"Non-standard font size found (expected {expected_size}pt)",
            "location": locations,
            "severity": "warning",
        })
    else:
        results.append({
            "check_type": "font_size",
            "passed": True,
            "message": f"Font size: {expected_size}pt OK",
            "location": "document",
            "severity": "info",
        })

    return results


def _check_spacing(doc: Document, profile: dict) -> list[dict]:
    expected = profile.get("line_spacing", 1.5)
    issues = 0

    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            continue
        pf = para.paragraph_format
        if pf.line_spacing is not None and isinstance(pf.line_spacing, (int, float)):
            if abs(pf.line_spacing - expected) > 0.1:
                issues += 1

    passed = issues == 0
    return [{
        "check_type": "line_spacing",
        "passed": passed,
        "message": f"Line spacing: {'OK' if passed else f'{issues} paragraphs with incorrect spacing'}",
        "location": "document",
        "severity": "info" if passed else "warning",
    }]
