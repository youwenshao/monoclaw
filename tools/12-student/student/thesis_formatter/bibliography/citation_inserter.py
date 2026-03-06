"""Insert formatted in-text citations and bibliography into a .docx document."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

from student.thesis_formatter.bibliography.bib_formatter import format_bibliography


def insert_citations(doc_path: str, references: list[dict], style: str) -> str:
    output_path = _output_path(doc_path, "_cited")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    ref_map = _build_cite_key_map(references)
    _replace_inline_citations(doc, ref_map, style)
    _append_bibliography(doc, references, style)
    doc.save(output_path)
    return output_path


def _build_cite_key_map(references: list[dict]) -> dict[str, dict]:
    mapping: dict[str, dict] = {}
    for ref in references:
        key = ref.get("cite_key", "")
        if key:
            mapping[key] = ref
    return mapping


def _replace_inline_citations(doc: Document, ref_map: dict[str, dict], style: str) -> None:
    cite_pattern = re.compile(r"\\cite\{([^}]+)\}")

    for para in doc.paragraphs:
        full_text = para.text
        if not cite_pattern.search(full_text):
            continue

        def _replacer(m: re.Match) -> str:
            keys = [k.strip() for k in m.group(1).split(",")]
            parts = []
            for key in keys:
                ref = ref_map.get(key)
                if ref:
                    author = ref.get("author", "Unknown").split(",")[0].strip()
                    year = ref.get("year", "n.d.")
                    parts.append(f"{author}, {year}")
                else:
                    parts.append(key)
            return f"({'; '.join(parts)})"

        new_text = cite_pattern.sub(_replacer, full_text)
        if new_text != full_text:
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)


def _append_bibliography(doc: Document, references: list[dict], style: str) -> None:
    doc.add_paragraph("Bibliography", style="Heading 1")
    formatted = format_bibliography(references, style)
    for entry in formatted:
        doc.add_paragraph(entry)


def _output_path(doc_path: str, suffix: str = "") -> str:
    p = Path(doc_path)
    return str(p.parent / f"{p.stem}{suffix}{p.suffix}")
