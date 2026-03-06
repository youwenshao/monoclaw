"""Apply page numbering: Roman for front matter, Arabic from Chapter 1."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def apply_page_numbering(doc_path: str) -> str:
    output_path = _output_path(doc_path, "_numbered")
    shutil.copy2(doc_path, output_path)

    doc = Document(output_path)
    chapter_idx = _find_chapter_section(doc)

    if chapter_idx is not None and chapter_idx > 0:
        _insert_section_break(doc, chapter_idx)

    for i, section in enumerate(doc.sections):
        fmt = "lowerRoman" if i == 0 else "decimal"
        _set_page_number_format(section, fmt)
        if i > 0:
            _restart_numbering(section)

    doc.save(output_path)
    return output_path


def _find_chapter_section(doc: Document) -> int | None:
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if text.startswith("chapter") or text.startswith("1 ") or text.startswith("1."):
            return i
    return None


def _insert_section_break(doc: Document, para_idx: int) -> None:
    para = doc.paragraphs[para_idx]
    pPr = para._element.get_or_add_pPr()
    sect_pr = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect_pr.append(sect_type)
    pPr.addprevious(sect_pr)


def _set_page_number_format(section, fmt: str) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)


def _restart_numbering(section) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def _output_path(doc_path: str, suffix: str = "") -> str:
    p = Path(doc_path)
    return str(p.parent / f"{p.stem}{suffix}{p.suffix}")
