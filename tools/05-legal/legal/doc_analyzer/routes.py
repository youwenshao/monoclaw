"""LegalDoc Analyzer FastAPI routes."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, Request, UploadFile, File
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates

from openclaw_shared.database import get_db
from openclaw_shared.mona_events import emit_event

router = APIRouter(prefix="/doc-analyzer", tags=["DocAnalyzer"])
templates = Jinja2Templates(directory="legal/dashboard/templates")


def _ctx(request: Request, **extra: object) -> dict:
    return {"request": request, "config": request.app.state.config, "active_tab": "doc-analyzer", **extra}


def _uploads_dir(request: Request) -> Path:
    workspace: Path = request.app.state.workspace
    d = workspace / "contracts" / "incoming"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _extract_text(file_path: Path) -> str:
    """Extract plain text from a contract file (PDF or DOCX)."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            pages = [page.get_text() for page in doc]
            doc.close()
            return "\n\n".join(pages)
        except ImportError:
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                return pdfminer_extract(str(file_path))
            except ImportError:
                return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(file_path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return file_path.read_text(encoding="utf-8", errors="replace")

    return file_path.read_text(encoding="utf-8", errors="replace")


def _count_pages(file_path: Path) -> int:
    """Estimate page count from a contract file."""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(file_path))
            count = doc.page_count
            doc.close()
            return count
        except ImportError:
            pass

    text = _extract_text(file_path)
    return max(1, len(text) // 3000)


# ── Page ───────────────────────────────────────────────────────────────────

@router.get("/", response_class=HTMLResponse)
async def doc_analyzer_page(request: Request) -> HTMLResponse:
    db = request.app.state.db_paths["doc_analyzer"]

    with get_db(db) as conn:
        contracts = [dict(r) for r in conn.execute(
            "SELECT * FROM contracts ORDER BY upload_date DESC"
        ).fetchall()]
        total_contracts = conn.execute("SELECT COUNT(*) FROM contracts").fetchone()[0]
        pending_count = conn.execute(
            "SELECT COUNT(*) FROM contracts WHERE analysis_status = 'pending'"
        ).fetchone()[0]
        completed_count = conn.execute(
            "SELECT COUNT(*) FROM contracts WHERE analysis_status = 'completed'"
        ).fetchone()[0]
        anomaly_count = conn.execute(
            "SELECT COUNT(*) FROM clauses WHERE anomaly_score > 0.5"
        ).fetchone()[0]
        avg_anomaly = conn.execute(
            "SELECT AVG(anomaly_score) FROM clauses WHERE anomaly_score IS NOT NULL"
        ).fetchone()[0] or 0.0

    return templates.TemplateResponse(
        "doc_analyzer/index.html",
        _ctx(
            request,
            contracts=contracts,
            total_contracts=total_contracts,
            pending_count=pending_count,
            completed_count=completed_count,
            anomaly_count=anomaly_count,
            avg_anomaly=round(avg_anomaly, 2),
        ),
    )


# ── Upload ─────────────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_contract(
    request: Request,
    file: UploadFile = File(...),
    contract_type: str = "other",
    language: str = "en",
) -> dict[str, Any]:
    db = request.app.state.db_paths["doc_analyzer"]
    uploads = _uploads_dir(request)

    filename = file.filename or "unnamed"
    file_path = uploads / filename
    content = await file.read()
    file_path.write_bytes(content)

    page_count = _count_pages(file_path)

    with get_db(db) as conn:
        cursor = conn.execute(
            """INSERT INTO contracts
               (filename, contract_type, language, analysis_status, page_count, file_path)
               VALUES (?,?,?,?,?,?)""",
            (filename, contract_type, language, "pending", page_count, str(file_path)),
        )
        contract_id = cursor.lastrowid

    emit_event(
        request.app.state.db_paths["mona_events"],
        event_type="action_started",
        tool_name="doc-analyzer",
        summary=f"Contract uploaded: {filename} (type: {contract_type})",
    )

    return {"contract_id": contract_id, "file_path": str(file_path), "status": "pending"}


# ── Analyze ────────────────────────────────────────────────────────────────

@router.post("/analyze/{contract_id}")
async def analyze_contract(request: Request, contract_id: int) -> dict[str, Any]:
    db = request.app.state.db_paths["doc_analyzer"]
    llm = request.app.state.llm

    with get_db(db) as conn:
        row = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = dict(row)
    file_path = Path(contract["file_path"]) if contract.get("file_path") else None

    if file_path and file_path.exists():
        text = _extract_text(file_path)
    else:
        with get_db(db) as conn:
            existing_clauses = conn.execute(
                "SELECT * FROM clauses WHERE contract_id = ?", (contract_id,)
            ).fetchall()
        if existing_clauses:
            text = "\n\n".join(dict(c)["text_content"] for c in existing_clauses)
        else:
            raise HTTPException(
                status_code=400,
                detail="Contract file not found and no existing clauses available",
            )

    with get_db(db) as conn:
        conn.execute(
            "UPDATE contracts SET analysis_status = 'analyzing' WHERE id = ?",
            (contract_id,),
        )

    from legal.doc_analyzer.clause_extractor import extract_clauses
    from legal.doc_analyzer.anomaly_detector import detect_anomalies

    extracted = await extract_clauses(text, contract["contract_type"], llm=llm)

    with get_db(db) as conn:
        ref_rows = conn.execute(
            "SELECT * FROM reference_clauses WHERE contract_type = ?",
            (contract["contract_type"],),
        ).fetchall()
    reference_clauses = [dict(r) for r in ref_rows]

    anomalies = await detect_anomalies(extracted, reference_clauses, llm=llm)

    anomaly_map = {a["clause_number"]: a for a in anomalies}

    with get_db(db) as conn:
        conn.execute("DELETE FROM clauses WHERE contract_id = ?", (contract_id,))

        for clause in extracted:
            cn = clause["clause_number"]
            anomaly = anomaly_map.get(cn, {})
            conn.execute(
                """INSERT INTO clauses
                   (contract_id, clause_number, clause_type, text_content,
                    anomaly_score, flag_reason, start_offset, end_offset)
                   VALUES (?,?,?,?,?,?,?,?)""",
                (
                    contract_id,
                    cn,
                    clause["clause_type"],
                    clause["text_content"],
                    anomaly.get("anomaly_score", 0.0),
                    anomaly.get("flag_reason"),
                    clause.get("start_offset"),
                    clause.get("end_offset"),
                ),
            )

        conn.execute(
            "UPDATE contracts SET analysis_status = 'completed' WHERE id = ?",
            (contract_id,),
        )

    compliance_result = None
    if contract["contract_type"] == "employment":
        from legal.doc_analyzer.employment_checker import check_cap57_compliance
        compliance_result = check_cap57_compliance(extracted)

    nda_result = None
    if contract["contract_type"] == "nda":
        from legal.doc_analyzer.nda_checker import check_nda_completeness
        nda_result = check_nda_completeness(extracted)

    flagged = [a for a in anomalies if a["anomaly_score"] > 0.5]
    emit_event(
        request.app.state.db_paths["mona_events"],
        event_type="action_completed",
        tool_name="doc-analyzer",
        summary=(
            f"Analysis complete for {contract['filename']}: "
            f"{len(extracted)} clauses, {len(flagged)} flagged"
        ),
        requires_human_action=len(flagged) > 0,
    )

    return {
        "contract_id": contract_id,
        "status": "completed",
        "clauses_extracted": len(extracted),
        "anomalies_flagged": len(flagged),
        "cap57_compliance": compliance_result,
        "nda_completeness": nda_result,
    }


# ── Contract detail ───────────────────────────────────────────────────────

@router.get("/contract/{contract_id}")
async def get_contract(request: Request, contract_id: int) -> dict[str, Any]:
    db = request.app.state.db_paths["doc_analyzer"]

    with get_db(db) as conn:
        row = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = dict(row)

    with get_db(db) as conn:
        clause_rows = conn.execute(
            "SELECT * FROM clauses WHERE contract_id = ? ORDER BY id",
            (contract_id,),
        ).fetchall()

    clauses = [dict(c) for c in clause_rows]

    compliance_result = None
    if contract["contract_type"] == "employment" and clauses:
        from legal.doc_analyzer.employment_checker import check_cap57_compliance
        compliance_result = check_cap57_compliance(clauses)

    nda_result = None
    if contract["contract_type"] == "nda" and clauses:
        from legal.doc_analyzer.nda_checker import check_nda_completeness
        nda_result = check_nda_completeness(clauses)

    return {
        "contract": contract,
        "clauses": clauses,
        "cap57_compliance": compliance_result,
        "nda_completeness": nda_result,
    }


# ── Compare ───────────────────────────────────────────────────────────────

@router.get("/compare")
async def compare_contracts(
    request: Request,
    contract_a: int | None = None,
    contract_b: int | None = None,
) -> dict[str, Any]:
    db = request.app.state.db_paths["doc_analyzer"]

    if contract_a is None or contract_b is None:
        with get_db(db) as conn:
            contracts = [dict(r) for r in conn.execute(
                "SELECT id, filename, contract_type, analysis_status FROM contracts "
                "WHERE analysis_status = 'completed' ORDER BY upload_date DESC"
            ).fetchall()]
        return {"contracts": contracts, "comparison": None}

    with get_db(db) as conn:
        row_a = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_a,)).fetchone()
        row_b = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_b,)).fetchone()

    if not row_a or not row_b:
        raise HTTPException(status_code=404, detail="One or both contracts not found")

    with get_db(db) as conn:
        clauses_a = [dict(r) for r in conn.execute(
            "SELECT * FROM clauses WHERE contract_id = ? ORDER BY id", (contract_a,)
        ).fetchall()]
        clauses_b = [dict(r) for r in conn.execute(
            "SELECT * FROM clauses WHERE contract_id = ? ORDER BY id", (contract_b,)
        ).fetchall()]

    types_a = {c["clause_type"] for c in clauses_a}
    types_b = {c["clause_type"] for c in clauses_b}
    all_types = sorted(types_a | types_b)

    comparison = []
    for clause_type in all_types:
        side_a = [c for c in clauses_a if c["clause_type"] == clause_type]
        side_b = [c for c in clauses_b if c["clause_type"] == clause_type]
        comparison.append({
            "clause_type": clause_type,
            "contract_a": side_a,
            "contract_b": side_b,
            "only_in_a": len(side_a) > 0 and len(side_b) == 0,
            "only_in_b": len(side_b) > 0 and len(side_a) == 0,
        })

    return {
        "contract_a": dict(row_a),
        "contract_b": dict(row_b),
        "comparison": comparison,
    }


# ── Export ─────────────────────────────────────────────────────────────────

@router.get("/export/{contract_id}")
async def export_contract(request: Request, contract_id: int) -> FileResponse:
    db = request.app.state.db_paths["doc_analyzer"]
    workspace: Path = request.app.state.workspace

    with get_db(db) as conn:
        row = conn.execute("SELECT * FROM contracts WHERE id = ?", (contract_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = dict(row)

    with get_db(db) as conn:
        clause_rows = conn.execute(
            "SELECT * FROM clauses WHERE contract_id = ? ORDER BY id",
            (contract_id,),
        ).fetchall()
    clauses = [dict(c) for c in clause_rows]

    export_dir = workspace / "contracts" / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_COLOR_INDEX

        doc = Document()
        doc.add_heading(f"Contract Analysis: {contract['filename']}", level=0)
        doc.add_paragraph(
            f"Type: {contract['contract_type']}  |  "
            f"Language: {contract['language']}  |  "
            f"Status: {contract['analysis_status']}"
        )
        doc.add_paragraph("")

        for clause in clauses:
            score = clause.get("anomaly_score", 0.0) or 0.0
            heading_text = f"Clause {clause['clause_number']} — {clause['clause_type']}"
            if score > 0.5:
                heading_text += f"  ⚠ (anomaly: {score:.0%})"

            heading = doc.add_heading(heading_text, level=2)
            if score > 0.7:
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

            para = doc.add_paragraph(clause.get("text_content", ""))
            if score > 0.5:
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            if clause.get("flag_reason"):
                flag_para = doc.add_paragraph()
                flag_run = flag_para.add_run(f"Flag: {clause['flag_reason']}")
                flag_run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
                flag_run.font.size = Pt(9)

            doc.add_paragraph("")

        output_name = f"analysis_{contract_id}_{contract['filename'].rsplit('.', 1)[0]}.docx"
        output_path = export_dir / output_name
        doc.save(str(output_path))

    except ImportError:
        output_name = f"analysis_{contract_id}_{contract['filename'].rsplit('.', 1)[0]}.json"
        output_path = export_dir / output_name
        export_data = {
            "contract": contract,
            "clauses": clauses,
        }
        output_path.write_text(json.dumps(export_data, indent=2, default=str))

    emit_event(
        request.app.state.db_paths["mona_events"],
        event_type="action_completed",
        tool_name="doc-analyzer",
        summary=f"Exported annotated analysis for {contract['filename']}",
    )

    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if output_path.suffix == ".docx"
        else "application/json"
    )
    return FileResponse(path=str(output_path), filename=output_name, media_type=media_type)


# ── Partials ──────────────────────────────────────────────────────────────

@router.get("/clause-view/partial", response_class=HTMLResponse)
async def clause_view_partial(request: Request, contract_id: int) -> HTMLResponse:
    db = request.app.state.db_paths["doc_analyzer"]

    with get_db(db) as conn:
        contract_row = conn.execute(
            "SELECT * FROM contracts WHERE id = ?", (contract_id,)
        ).fetchone()
        clause_rows = conn.execute(
            "SELECT * FROM clauses WHERE contract_id = ? ORDER BY id",
            (contract_id,),
        ).fetchall()

    if not contract_row:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = dict(contract_row)
    clauses = [dict(c) for c in clause_rows]

    compliance_result = None
    if contract["contract_type"] == "employment" and clauses:
        from legal.doc_analyzer.employment_checker import check_cap57_compliance
        compliance_result = check_cap57_compliance(clauses)

    nda_result = None
    if contract["contract_type"] == "nda" and clauses:
        from legal.doc_analyzer.nda_checker import check_nda_completeness
        nda_result = check_nda_completeness(clauses)

    return templates.TemplateResponse(
        "doc_analyzer/partials/clause_view.html",
        {
            "request": request,
            "contract": contract,
            "clauses": clauses,
            "cap57_compliance": compliance_result,
            "nda_completeness": nda_result,
        },
    )


@router.get("/anomaly-detail/partial", response_class=HTMLResponse)
async def anomaly_detail_partial(request: Request, clause_id: int) -> HTMLResponse:
    db = request.app.state.db_paths["doc_analyzer"]

    with get_db(db) as conn:
        clause_row = conn.execute(
            "SELECT * FROM clauses WHERE id = ?", (clause_id,)
        ).fetchone()

    if not clause_row:
        raise HTTPException(status_code=404, detail="Clause not found")

    clause = dict(clause_row)

    with get_db(db) as conn:
        ref_rows = conn.execute(
            "SELECT * FROM reference_clauses WHERE clause_type = ?",
            (clause["clause_type"],),
        ).fetchall()
    references = [dict(r) for r in ref_rows]

    return templates.TemplateResponse(
        "doc_analyzer/partials/anomaly_detail.html",
        {
            "request": request,
            "clause": clause,
            "references": references,
        },
    )
