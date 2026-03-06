"""Generate draft engagement letter documents using python-docx."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger("openclaw.legal.intake_bot.engagement_letter")


def generate_engagement_letter(
    client_data: dict[str, Any],
    matter_data: dict[str, Any],
    firm_config: dict[str, Any],
    output_path: str | Path,
) -> Path:
    """Generate a draft engagement letter as a .docx file.

    Includes: firm letterhead, client address, RE line, scope of retainer,
    fee arrangement (hourly/fixed/conditional), terms, and signature block.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    firm_name = firm_config.get("firm_name", "")
    hkls_reg = firm_config.get("hkls_registration", "")
    office_address = firm_config.get("office_address", "")

    p_firm = doc.add_paragraph()
    p_firm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_firm.add_run(firm_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(26, 58, 92)

    if hkls_reg:
        p_reg = doc.add_paragraph()
        p_reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_reg.add_run(f"HKLS Registration: {hkls_reg}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    if office_address:
        p_addr = doc.add_paragraph()
        p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_addr.add_run(office_address)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    _add_horizontal_rule(doc)

    today_str = datetime.now().strftime("%d %B %Y")
    doc.add_paragraph(today_str)
    doc.add_paragraph()

    client_name_en = client_data.get("name_en", "")
    client_name_tc = client_data.get("name_tc", "")
    client_email = client_data.get("email", "")
    client_phone = client_data.get("phone", "")

    display_name = client_name_en
    if client_name_tc:
        display_name = f"{client_name_en} ({client_name_tc})"

    p_to = doc.add_paragraph()
    p_to.add_run(display_name).bold = True
    if client_email:
        doc.add_paragraph(f"Email: {client_email}")
    if client_phone:
        doc.add_paragraph(f"Phone: {client_phone}")

    doc.add_paragraph()

    p_salutation = doc.add_paragraph(f"Dear {client_name_en or 'Client'},")
    doc.add_paragraph()

    matter_type = matter_data.get("matter_type", "general")
    if matter_type:
        matter_type = matter_type.replace("_", " ").title()
    description = matter_data.get("description", "")
    adverse_party = matter_data.get("adverse_party_name", "")

    re_line = f"RE: {matter_type}"
    if adverse_party:
        re_line += f" — {client_name_en} v. {adverse_party}"
    p_re = doc.add_paragraph()
    run = p_re.add_run(re_line)
    run.bold = True
    run.underline = True

    doc.add_paragraph()

    doc.add_paragraph(
        f"We are pleased to confirm that {firm_name} has agreed to act on your behalf "
        f"in the above matter. This letter sets out the terms on which we will provide "
        f"our legal services."
    )

    _add_section_heading(doc, "1. Scope of Retainer")
    scope_text = (
        f"We are retained to advise and represent you in relation to: {description or matter_type}."
    )
    if adverse_party:
        scope_text += f" The adverse party in this matter is {adverse_party}."
    scope_text += (
        " Our retainer is limited to this matter and does not extend to any other "
        "legal issues unless separately agreed in writing."
    )
    doc.add_paragraph(scope_text)

    _add_section_heading(doc, "2. Responsible Solicitor")
    solicitor = matter_data.get("assigned_solicitor", "a designated solicitor")
    doc.add_paragraph(
        f"The solicitor with primary responsibility for this matter will be {solicitor}. "
        f"Other members of the firm may assist from time to time."
    )

    _add_section_heading(doc, "3. Fee Arrangement")

    fee_type = firm_config.get("default_fee_type", "hourly")
    hourly_rate = firm_config.get("hourly_rate", "HK$3,000 – HK$6,000")

    if fee_type == "hourly":
        doc.add_paragraph(
            f"Our fees will be charged on a time-spent basis at the following hourly rates: "
            f"{hourly_rate} per hour (depending on the seniority of the solicitor involved). "
            f"Time is recorded in six-minute units. Disbursements (e.g. court filing fees, "
            f"counsel's fees, search fees) will be charged separately at cost."
        )
    elif fee_type == "fixed":
        fixed_amount = firm_config.get("fixed_fee_amount", "to be agreed")
        doc.add_paragraph(
            f"A fixed fee of {fixed_amount} has been agreed for this matter. "
            f"This fee covers the scope of work described above. Any work outside "
            f"this scope will be subject to separate agreement. Disbursements will "
            f"be charged separately at cost."
        )
    elif fee_type == "conditional":
        doc.add_paragraph(
            "Our fees in this matter will be charged on a conditional basis as permitted "
            "under the Legal Practitioners Ordinance (Cap. 159). Details of the conditional "
            "fee arrangement are set out in the separate Conditional Fee Agreement."
        )

    doc.add_paragraph(
        "We will issue interim bills at monthly or quarterly intervals. Payment terms "
        "are 14 days from the date of invoice."
    )

    _add_section_heading(doc, "4. Client Monies")
    doc.add_paragraph(
        "Any monies received by us on your behalf will be held in our client account "
        "in accordance with the Solicitors' Accounts Rules. We may request money on "
        "account of fees and disbursements from time to time."
    )

    _add_section_heading(doc, "5. Confidentiality")
    doc.add_paragraph(
        "All information you provide to us will be treated as strictly confidential "
        "and subject to legal professional privilege, except where disclosure is "
        "required by law or with your consent."
    )

    _add_section_heading(doc, "6. Termination")
    doc.add_paragraph(
        "Either party may terminate this retainer by giving reasonable written notice. "
        "In the event of termination, you will be responsible for fees and disbursements "
        "incurred up to the date of termination. We will provide all necessary assistance "
        "to facilitate the transfer of the matter to another firm."
    )

    _add_section_heading(doc, "7. Complaints")
    doc.add_paragraph(
        "If you have any concerns about our services, please raise them with the "
        "responsible solicitor in the first instance. If the matter is not resolved "
        "to your satisfaction, you may refer the complaint to The Law Society of "
        "Hong Kong."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Please confirm your acceptance of these terms by signing and returning "
        "a copy of this letter."
    )

    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()

    p_sig = doc.add_paragraph()
    p_sig.add_run("_" * 40)
    doc.add_paragraph(f"For and on behalf of {firm_name}")
    if solicitor:
        doc.add_paragraph(solicitor)

    doc.add_paragraph()
    doc.add_paragraph()
    _add_horizontal_rule(doc)

    doc.add_paragraph("ACKNOWLEDGED AND AGREED:")
    doc.add_paragraph()
    p_client_sig = doc.add_paragraph()
    p_client_sig.add_run("_" * 40)
    doc.add_paragraph(f"Name: {display_name}")
    p_date_sig = doc.add_paragraph()
    p_date_sig.add_run("Date: _" * 10)

    doc.save(str(out))
    logger.info("Engagement letter generated: %s", out)
    return out


def _add_section_heading(doc: Any, text: str) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(26, 58, 92)


def _add_horizontal_rule(doc: Any) -> None:
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("—" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(192, 192, 192)
