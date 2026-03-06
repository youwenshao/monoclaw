"""Generate a Provisional Tenancy Agreement (臨時租約) in DOCX format.

A shorter, binding document typically signed at the point of agreement
before the formal tenancy agreement is executed.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def generate_provisional_agreement(tenancy_data: dict, output_dir: Path) -> Path:
    """Create a provisional tenancy agreement and return the file path.

    ``tenancy_data`` should include::

        landlord_name, tenant_name, property_address,
        monthly_rent, term_months, start_date, end_date

    Optional: ``property_address_zh``, ``landlord_hkid``, ``tenant_hkid``,
    ``deposit_months`` (default 2), ``commission_split``, ``handover_date``,
    ``special_conditions``, ``landlord_agent``, ``tenant_agent``.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    landlord = tenancy_data["landlord_name"]
    tenant = tenancy_data["tenant_name"]
    address = tenancy_data["property_address"]
    address_zh = tenancy_data.get("property_address_zh", address)
    rent = tenancy_data["monthly_rent"]
    deposit_months = tenancy_data.get("deposit_months", 2)
    deposit = tenancy_data.get("deposit_amount", rent * deposit_months)
    term = tenancy_data["term_months"]
    start = tenancy_data["start_date"]
    end = tenancy_data["end_date"]
    handover = tenancy_data.get("handover_date", start)
    commission_split = tenancy_data.get("commission_split", "50/50")
    landlord_hkid = tenancy_data.get("landlord_hkid", "")
    tenant_hkid = tenancy_data.get("tenant_hkid", "")
    landlord_agent = tenancy_data.get("landlord_agent", "")
    tenant_agent = tenancy_data.get("tenant_agent", "")

    # ── Title ──────────────────────────────────────────────────────────
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title.add_run("PROVISIONAL TENANCY AGREEMENT")
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    title.add_run("\n")
    r2 = title.add_run("臨時租賃協議")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        f"Date 日期: {datetime.now().strftime('%d %B %Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Summary Table ──────────────────────────────────────────────────
    table = doc.add_table(rows=10, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("Landlord 業主", f"{landlord}" + (f"  (HKID: {landlord_hkid})" if landlord_hkid else "")),
        ("Tenant 租客", f"{tenant}" + (f"  (HKID: {tenant_hkid})" if tenant_hkid else "")),
        ("Premises 物業地址", f"{address}\n{address_zh}"),
        ("Term 租期", f"{term} months 個月 ({start} to 至 {end})"),
        ("Monthly Rent 每月租金", f"HK${rent:,}"),
        ("Deposit 按金", f"HK${deposit:,} ({deposit_months} months' rent 個月租金)"),
        ("Handover Date 交樓日期", str(handover)),
        ("Commission Split 佣金分配", commission_split),
        ("Landlord Agent 業主代理", landlord_agent or "N/A"),
        ("Tenant Agent 租客代理", tenant_agent or "N/A"),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for cell in (table.cell(i, 0), table.cell(i, 1)):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Key Terms ──────────────────────────────────────────────────────
    h = doc.add_heading("KEY TERMS 主要條款", level=2)

    terms = [
        (
            f"Upon signing this Provisional Agreement, the Tenant shall pay a "
            f"deposit of HK${deposit:,} (equivalent to {deposit_months} months' "
            f"rent) to the Landlord or the Landlord's agent.",
            f"簽署本臨時協議時，租客須向業主或業主代理繳付按金港幣 ${deposit:,} 元"
            f"（相等於 {deposit_months} 個月租金）。",
        ),
        (
            f"The formal Tenancy Agreement shall be executed within 14 days of "
            f"the date hereof. In the event the Tenant fails to execute the "
            f"formal Agreement, the deposit shall be forfeited.",
            "正式租賃協議須於本協議簽署日期起十四日內簽署。如租客未能簽署正式協議，已繳按金將被沒收。",
        ),
        (
            f"The Landlord shall deliver vacant possession of the Premises to "
            f"the Tenant on {handover} in good and clean condition.",
            f"業主須於 {handover} 將物業以良好清潔狀態交付予租客作空置管有。",
        ),
        (
            f"Estate agent commission shall be split {commission_split} between "
            f"the Landlord and the Tenant (each party paying their respective "
            f"agent's commission equivalent to half month's rent).",
            f"地產代理佣金由業主及租客按 {commission_split} 分攤（各方支付相等於半個月租金之佣金予其代理）。",
        ),
        (
            "Both parties agree to bear the stamp duty equally.",
            "雙方同意平均分擔印花稅。",
        ),
    ]

    for i, (en, zh) in enumerate(terms, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(en).font.size = Pt(10)
        p.add_run(f"\n    {zh}").font.size = Pt(10)

    # ── Special Conditions ─────────────────────────────────────────────
    special = tenancy_data.get("special_conditions")
    if special:
        doc.add_heading("SPECIAL CONDITIONS 特別條款", level=2)
        if isinstance(special, list):
            for i, cond in enumerate(special, 1):
                doc.add_paragraph(f"S{i}. {cond}")
        else:
            doc.add_paragraph(str(special))

    # ── Acknowledgement ────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Both parties confirm that they have read, understood, and agree to "
        "be bound by the terms of this Provisional Tenancy Agreement."
    ).font.size = Pt(10)
    p.add_run(
        "\n雙方確認已閱讀、明瞭並同意受本臨時租賃協議條款所約束。"
    ).font.size = Pt(10)

    # ── Signature Blocks ───────────────────────────────────────────────
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = True
    sig.cell(0, 0).text = (
        "____________________________\n"
        f"Landlord 業主: {landlord}\n"
        f"Date 日期: ___________________\n\n"
        "____________________________\n"
        f"Landlord Agent 業主代理: {landlord_agent or '_______________'}"
    )
    sig.cell(0, 1).text = (
        "____________________________\n"
        f"Tenant 租客: {tenant}\n"
        f"Date 日期: ___________________\n\n"
        "____________________________\n"
        f"Tenant Agent 租客代理: {tenant_agent or '_______________'}"
    )

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = (
        f"provisional_agreement_{tenancy_data.get('id', 'draft')}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    path = output_dir / filename
    doc.save(str(path))
    return path
