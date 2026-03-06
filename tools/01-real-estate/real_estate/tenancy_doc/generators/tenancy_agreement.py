"""Generate a SAR-standard tenancy agreement in DOCX format.

Produces a bilingual (English / Traditional Chinese) document covering
all standard Hong Kong residential tenancy clauses.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor


def _heading(doc: Document, en: str, zh: str, level: int = 2) -> None:
    h = doc.add_heading(level=level)
    run_en = h.add_run(en)
    run_en.font.size = Pt(13 if level == 2 else 16)
    h.add_run("  ")
    run_zh = h.add_run(zh)
    run_zh.font.size = Pt(13 if level == 2 else 16)


def _para(doc: Document, en: str, zh: str) -> None:
    p = doc.add_paragraph()
    p.add_run(en).font.size = Pt(10)
    p.add_run("\n")
    run_zh = p.add_run(zh)
    run_zh.font.size = Pt(10)


def _bold_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)


def generate_tenancy_agreement(tenancy_data: dict, output_dir: Path) -> Path:
    """Create a formal tenancy agreement DOCX and return the file path.

    ``tenancy_data`` should contain at minimum::

        landlord_name, tenant_name, property_address,
        monthly_rent, deposit_amount, term_months,
        start_date, end_date

    Optional fields: ``property_address_zh``, ``landlord_hkid``,
    ``tenant_hkid``, ``break_clause_date``, ``special_conditions``.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    landlord = tenancy_data["landlord_name"]
    tenant = tenancy_data["tenant_name"]
    address = tenancy_data["property_address"]
    address_zh = tenancy_data.get("property_address_zh", address)
    rent = tenancy_data["monthly_rent"]
    deposit = tenancy_data.get("deposit_amount", rent * 2)
    term = tenancy_data["term_months"]
    start = tenancy_data["start_date"]
    end = tenancy_data["end_date"]
    break_date = tenancy_data.get("break_clause_date")
    landlord_hkid = tenancy_data.get("landlord_hkid", "")
    tenant_hkid = tenancy_data.get("tenant_hkid", "")

    # ── Title ──────────────────────────────────────────────────────────
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title.add_run("TENANCY AGREEMENT")
    r1.font.size = Pt(18)
    r1.font.color.rgb = RGBColor(0, 0, 0)
    title.add_run("\n")
    r2 = title.add_run("租賃協議")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        f"Date 日期: {datetime.now().strftime('%d %B %Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 1. Parties ─────────────────────────────────────────────────────
    _heading(doc, "1. PARTIES", "第一條  訂約方")
    _para(
        doc,
        f"LANDLORD 業主: {landlord}"
        + (f"  (HKID: {landlord_hkid})" if landlord_hkid else ""),
        f"業主: {landlord}",
    )
    _para(
        doc,
        f"TENANT 租客: {tenant}"
        + (f"  (HKID: {tenant_hkid})" if tenant_hkid else ""),
        f"租客: {tenant}",
    )

    # ── 2. Premises ────────────────────────────────────────────────────
    _heading(doc, "2. PREMISES", "第二條  租賃物業")
    _para(
        doc,
        f"The Landlord agrees to let and the Tenant agrees to take the premises situated at:\n{address}",
        f"業主同意將以下物業出租予租客，租客同意承租：\n{address_zh}",
    )

    # ── 3. Term ────────────────────────────────────────────────────────
    _heading(doc, "3. TERM", "第三條  租期")
    _para(
        doc,
        f"The tenancy shall be for a fixed term of {term} months, "
        f"commencing on {start} and expiring on {end} (both dates inclusive).",
        f"租期為固定 {term} 個月，由 {start} 起至 {end} 止（包括首尾兩天）。",
    )

    # ── 4. Rent ────────────────────────────────────────────────────────
    _heading(doc, "4. RENT", "第四條  租金")
    _para(
        doc,
        f"The monthly rent shall be HK${rent:,}, payable in advance on or before "
        "the first day of each calendar month by bank transfer or cheque.",
        f"每月租金為港幣 ${rent:,} 元，須於每月首日或之前以銀行轉帳或支票預繳。",
    )

    # ── 5. Deposit ─────────────────────────────────────────────────────
    _heading(doc, "5. DEPOSIT", "第五條  按金")
    _para(
        doc,
        f"The Tenant shall pay a deposit of HK${deposit:,} (equivalent to "
        f"{deposit // rent if rent else 0} months' rent) upon signing this Agreement. "
        "The deposit shall be refunded without interest within 30 days after "
        "the expiry of the tenancy, subject to deduction for outstanding rent, "
        "utilities, or damages beyond fair wear and tear.",
        f"租客須於簽署本協議時繳付按金港幣 ${deposit:,} 元"
        f"（相等於 {deposit // rent if rent else 0} 個月租金）。"
        "按金將於租約届滿後三十日內退還（不計利息），惟業主有權扣除"
        "欠繳租金、水電費或超出正常損耗之維修費用。",
    )

    # ── 6. Break Clause ────────────────────────────────────────────────
    if break_date:
        _heading(doc, "6. BREAK CLAUSE", "第六條  提前終止條款")
        _para(
            doc,
            f"Either party may terminate this Agreement on {break_date} "
            "by giving not less than two (2) months' prior written notice "
            "to the other party.",
            f"任何一方均可於 {break_date} 提前終止本協議，惟須給予另一方不少於兩個月的書面通知。",
        )
        clause_num = 7
    else:
        clause_num = 6

    # ── Government Rent & Rates ────────────────────────────────────────
    _heading(doc, f"{clause_num}. GOVERNMENT RENT AND RATES", f"第{clause_num}條  地租及差餉")
    _para(
        doc,
        "The Landlord shall be responsible for payment of Government Rent. "
        "The Tenant shall be responsible for payment of Rates, "
        "Management Fees, and all utility charges during the tenancy period.",
        "業主負責繳納地租。租客負責繳納差餉、管理費及租賃期間之所有水電煤費用。",
    )
    clause_num += 1

    # ── Maintenance ────────────────────────────────────────────────────
    _heading(doc, f"{clause_num}. REPAIR AND MAINTENANCE", f"第{clause_num}條  維修及保養")
    _para(
        doc,
        "The Tenant shall keep the interior of the Premises in good and "
        "tenantable condition (fair wear and tear excepted) and shall be "
        "responsible for minor repairs. The Landlord shall be responsible "
        "for structural repairs and repair of fixtures and installations "
        "installed by the Landlord.",
        "租客須保持物業內部良好及適宜居住之狀況（正常損耗除外），並負責日常小型維修。"
        "業主負責結構性維修及由業主安裝之裝置設備之維修。",
    )
    clause_num += 1

    # ── Use of Premises ────────────────────────────────────────────────
    _heading(doc, f"{clause_num}. USE OF PREMISES", f"第{clause_num}條  物業用途")
    _para(
        doc,
        "The Premises shall be used solely for residential purposes and shall "
        "not be used for any illegal or immoral purposes or in any way that "
        "may be a nuisance or annoyance to the occupiers of neighbouring premises.",
        "物業僅供住宅用途，不得用作任何非法或不道德之目的，亦不得以任何方式對鄰近物業之住戶構成滋擾。",
    )
    clause_num += 1

    # ── Assignment / Subletting ────────────────────────────────────────
    _heading(doc, f"{clause_num}. ASSIGNMENT AND SUBLETTING", f"第{clause_num}條  轉讓及分租")
    _para(
        doc,
        "The Tenant shall not assign, sublet, or part with possession of "
        "the Premises or any part thereof without the prior written consent "
        "of the Landlord.",
        "未經業主事先書面同意，租客不得轉讓、分租或放棄物業或其任何部分之佔有。",
    )
    clause_num += 1

    # ── Termination ────────────────────────────────────────────────────
    _heading(doc, f"{clause_num}. TERMINATION", f"第{clause_num}條  終止")
    _para(
        doc,
        "If the Tenant fails to pay rent for 15 days after the due date, or "
        "commits a material breach of any term of this Agreement, the Landlord "
        "may re-enter the Premises and this Agreement shall thereupon determine, "
        "without prejudice to any right of action the Landlord may have against "
        "the Tenant.",
        "如租客逾期十五日未繳租金，或嚴重違反本協議任何條款，業主有權重新進入物業，"
        "本協議即告終止，惟不影響業主對租客可能享有之任何訴訟權利。",
    )
    clause_num += 1

    # ── Stamp Duty ─────────────────────────────────────────────────────
    _heading(doc, f"{clause_num}. STAMP DUTY", f"第{clause_num}條  印花稅")
    _para(
        doc,
        "Stamp duty on this Agreement shall be borne equally by the Landlord "
        "and the Tenant. The Agreement shall be submitted for stamping within "
        "30 days of execution.",
        "本協議之印花稅由業主及租客各付一半。本協議須於簽署後三十日內提交蓋印。",
    )
    clause_num += 1

    # ── Governing Law ──────────────────────────────────────────────────
    _heading(doc, f"{clause_num}. GOVERNING LAW", f"第{clause_num}條  適用法律")
    _para(
        doc,
        "This Agreement shall be governed by and construed in accordance with "
        "the laws of the Hong Kong Special Administrative Region.",
        "本協議受香港特別行政區法律管轄並據其解釋。",
    )
    clause_num += 1

    # ── Special Conditions ─────────────────────────────────────────────
    special = tenancy_data.get("special_conditions")
    if special:
        _heading(doc, f"{clause_num}. SPECIAL CONDITIONS", f"第{clause_num}條  特別條款")
        if isinstance(special, list):
            for i, cond in enumerate(special, 1):
                doc.add_paragraph(f"{clause_num}.{i}  {cond}", style="List Number")
        else:
            doc.add_paragraph(str(special))
        clause_num += 1

    # ── Signature Blocks ───────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_table(rows=2, cols=2)
    sig.autofit = True
    sig.cell(0, 0).text = (
        "____________________________\n"
        f"Landlord 業主: {landlord}\n"
        f"Date 日期: ___________________"
    )
    sig.cell(0, 1).text = (
        "____________________________\n"
        f"Tenant 租客: {tenant}\n"
        f"Date 日期: ___________________"
    )
    sig.cell(1, 0).text = (
        "____________________________\n"
        "Witness 見證人:\n"
        "HKID: ______________________"
    )
    sig.cell(1, 1).text = (
        "____________________________\n"
        "Witness 見證人:\n"
        "HKID: ______________________"
    )

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = (
        f"tenancy_agreement_{tenancy_data.get('id', 'draft')}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    path = output_dir / filename
    doc.save(str(path))
    return path
